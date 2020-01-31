from pulp import LpProblem, LpMaximize, LpVariable, LpStatus, value, lpSum, LpStatusOptimal
import numpy as np
from docx.enum.text import WD_ALIGN_PARAGRAPH
from report import Report

def plan(products, resourceCount, periodsCount, fonds, resourceConsumption,
         mvp):
    doc = Report()
    problem = LpProblem("Calendar_planning", LpMaximize)
    doc.d.add_heading('Исходные данные', 2)
    doc.add_paragraph('Количество изделий: %d.' % len(products))
    doc.add_paragraph('Количество ресурсов: %d.' % resourceCount)
    doc.add_paragraph('Количество интервалов планирования: %d.' % periodsCount)
    # Годовой план
    doc.add_paragraph('Годовой план', WD_ALIGN_PARAGRAPH.CENTER)
    table = doc.d.add_table(1, 3)
    table.autofit = True
    table_head = table.rows[0].cells
    table_head[0].text = '№'
    table_head[1].text = 'Выпуск в год'
    table_head[2].text = 'Приоритет'
    for p in products:
        row = table.add_row().cells
        row[0].text = str(p['id'])
        row[1].text = str(p['annual'])
        row[2].text = str(p['priority'])

    # Фонды ресурсов
    doc.add_paragraph('Фонды ресурсов', WD_ALIGN_PARAGRAPH.CENTER)
    table = doc.d.add_table(2, resourceCount + 1)
    table.autofit = True
    table_head = table.rows[0].cells
    table_head[1].merge(table_head[resourceCount]).text = 'Ресурс'
    table_head = table.rows[1].cells 
    table_head[0].text = 'Период'
    for i in range(1, resourceCount + 1):
        table_head[i].text = str(i)
    for i, f in enumerate(fonds):
        row = table.add_row().cells
        row[0].text = str(i + 1)
        for j, v in enumerate(f):
            row[j + 1].text = str(v)

    # Нормы расхода ресурсов
    doc.add_paragraph('Нормы расхода ресурсов на единицу продукции', WD_ALIGN_PARAGRAPH.CENTER)
    table = doc.d.add_table(2, resourceCount + 1)
    table.autofit = True
    table_head = table.rows[0].cells
    table_head[1].merge(table_head[resourceCount]).text = 'Ресурс'
    table_head = table.rows[1].cells 
    table_head[0].text = '№ изделия'
    for i in range(1, resourceCount + 1):
        table_head[i].text = str(i)
    for i, rc in enumerate(resourceConsumption):
        row = table.add_row().cells
        row[0].text = str(i + 1)
        for j, v in enumerate(rc):
            row[j + 1].text = str(v)

    # Минимально необходимые партии изделий
    doc.add_paragraph('Минимально необходимые партии изделий', WD_ALIGN_PARAGRAPH.CENTER)
    table = doc.d.add_table(2, len(products) + 1)
    table.autofit = True
    table_head = table.rows[0].cells
    table_head[1].merge(table_head[len(products)]).text = '№ изделия'
    table_head = table.rows[1].cells 
    table_head[0].text = 'Период'
    for i in range(1, len(products) + 1):
        table_head[i].text = str(i)
    for i, m in enumerate(mvp):
        row = table.add_row().cells
        row[0].text = str(i + 1)
        for j, v in enumerate(m):
            row[j + 1].text = str(v)

    z = 0
    periodPlan = []
    for period in range(1, periodsCount + 1):
        periodPlan.append([])
        for product in products:
            var_product = LpVariable(
                f"x_{product['id']}_{period}",
                lowBound=mvp[period - 1][product['id']],
                cat='Integer')
            z += (1.0 / (product['priority'] * period)) * var_product
            periodPlan[period - 1].append(var_product)
        for resource in range(resourceCount):
            problem += lpSum([
                rc[resource] * periodPlan[period - 1][i]
                for i, rc in enumerate(resourceConsumption)
            ]) <= fonds[period - 1][resource], ""

    for product in products:
        problem += lpSum([p[product['id']]
                          for p in periodPlan]) == product['annual'], ""

    problem += z
    problem.solve()
    if problem.status == LpStatusOptimal:
        doc.add_paragraph('Полученное решение', WD_ALIGN_PARAGRAPH.CENTER)
        table = doc.d.add_table(2, periodsCount + 1)
        table_head = table.rows[0].cells
        table_head[1].merge(table_head[periodsCount]).text = 'Период'
        table_head = table.rows[1].cells
        table_head[0].text = 'Изделие'
        for i in range(1, periodsCount + 1):
            table_head[i].text = str(i)
        result = {
            'products': [],
            'report': ''
        }
        row = None
        for v in problem.variables():
            if v.name == '__dummy': continue
            *_, productId, period = v.name.split('_')
            productId = int(productId)
            period = int(period)
            if len(result['products']) <= productId:
                row = table.add_row().cells
                result['products'].append([])
            row[0].text = str(productId + 1)
            row[period].text = str(v.varValue)
            result['products'][productId].insert(period, v.varValue)
            # result.append({
            #     'name': v.name,
            #     'value': v.varValue
            # })
        file_path = doc.save('docs')
        result['report'] = f'https://toau1.herokuapp.com/report/{file_path}'
        #result['report'] = f'http://localhost:8000/report/{file_path}'
        return result
    else:
        return "Sosi hui"

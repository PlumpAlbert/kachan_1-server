from os import path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

class Report:
    def __init__(self):
        self.d = Document()
        # write title page
        t = self.d.add_heading('Распределение производственной программы при длительности '\
                  'производственного цикла меньше интервала планирования', 1)


    def add_paragraph(self, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        p = self.d.add_paragraph(text)
        f = p.paragraph_format
        f.alignment = align
        f.first_line_indent = Cm(1.25)
        return p


    def save(self, base_dir):
        file_path = path.join(base_dir, str(datetime.now()) + '.docx')
        self.d.save(file_path)
        return file_path
